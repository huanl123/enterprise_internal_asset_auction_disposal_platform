import os,re
from docx import Document

root=r"e:\\Study\\code\\Java\\enterprise_internal_asset_auction_disposal_platform\\毕设论文要求"
files=['2215304659 庄采润 论文.docx','2215304659 庄采润 论文_打印版.docx']
new_ref20='[20] OWASP Foundation. OWASP API Security Top 10 - 2023[EB/OL]. https://owasp.org/API-Security/editions/2023/en/0x11-t10/, 2026-03-05.'

for fn in files:
    path=os.path.join(root,fn)
    d=Document(path)

    # replace [20] reference line
    replaced=False
    for p in d.paragraphs:
        t=p.text.strip()
        if t=='[20]' or re.match(r'^\[20\]\s*$',t):
            p._p.clear_content()
            p.add_run(new_ref20)
            replaced=True
            break
        if t.startswith('[20]') and len(t)<20:
            p._p.clear_content()
            p.add_run(new_ref20)
            replaced=True
            break

    # ensure explicit citation [20] in body security paragraph
    injected=False
    for p in d.paragraphs:
        t=p.text.strip()
        if '安全测试重点检查未认证访问' in t and '[20]' not in t:
            p._p.clear_content()
            p.add_run(t + ' 同时，接口级安全风险评估可参考OWASP API Security Top 10（2023）方法进行基线核查[20]。')
            injected=True
            break

    d.save(path)
    print('UPDATED',fn,'ref20_replaced',replaced,'citation_injected',injected)
