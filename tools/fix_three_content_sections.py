import shutil
from pathlib import Path

from docx import Document


PAPER = Path(r"E:\Study\code\Java\enterprise_internal_asset_auction_disposal_platform\images\2215304659 庄采润 论文.docx")
BACKUP = Path(r"E:\Study\code\Java\enterprise_internal_asset_auction_disposal_platform\images\2215304659 庄采润 论文.bak_before_fix_three_sections.docx")


NEW_KEYWORDS = "Keywords  Fixed Asset Disposal  Internal Auction  Closed-Loop Workflow  Access Control  Statistical Analysis"

USECASE_P1 = (
    "以资产建档用例为例，资产专员登录系统后在资产管理页面录入待处置资产的基础信息，"
    "包括资产名称、类别、规格数量、所属部门、存放地点、购置时间、原值、残值、折旧规则及资产图片等内容。"
    "系统对必填项和数据格式进行校验，校验通过后生成资产档案，并将资产状态置为“待审核/待定价审核”，"
    "进入后续审核流程。"
)

USECASE_P2 = (
    "若财务专员在定价审核环节未通过该资产申请，系统将记录审核意见，并把资产状态退回至待补充资料或待重新提交状态；"
    "资产专员根据反馈补充或修改信息后，可再次提交进入审核环节。该用例说明了系统在资产建档、状态流转与异常回退方面的基本处理逻辑，"
    "其业务活动流程如图3.2所示。"
)

KEY_IMPL_P1 = (
    "为说明系统关键业务的实现方式，本文选取财务审核模块中的付款审核逻辑作为代表进行说明。"
    "该功能由 Spring Boot 服务层方法实现，并通过 @Transactional 注解保证交易单、资产状态、"
    "用户竞拍资格及操作记录在同一事务中完成更新，从而避免审核过程中出现部分成功、部分失败的数据不一致问题。"
    "方法执行时，系统首先依据交易单编号查询记录，并校验交易确认状态是否为 confirmed、付款状态是否为 pending，"
    "仅允许处于合法流程节点的交易进入付款审核环节。"
)

KEY_IMPL_P2 = (
    "当审核通过时，系统会校验并保存付款凭证，更新付款状态、付款时间和审核备注，同时写入操作历史；"
    "当审核不通过时，系统要求财务专员填写原因，将交易单状态更新为 rejected，并将资产状态回退为“待拍卖”，"
    "同时对违约中标人设置三个月竞拍限制。通过上述处理，系统实现了付款审核、异常回滚、违约处罚和业务留痕的统一控制，"
    "保证了资产处置流程的完整性、可追溯性和严谨性。"
)


def remove_paragraph(paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    parent.remove(el)


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(PAPER, BACKUP)

    doc = Document(str(PAPER))
    paras = doc.paragraphs

    # 1. English keywords
    kw_idx = next(i for i, p in enumerate(paras) if p.text.strip().startswith("Keywords"))
    paras[kw_idx].text = NEW_KEYWORDS

    # 2. Section 3.2.2 use-case description
    usecase_idx = next(i for i, p in enumerate(paras) if p.text.strip() == "3.2.2  主要用例说明")
    fig32_idx = next(i for i, p in enumerate(paras) if p.text.strip().startswith("图3.2"))
    paras[usecase_idx + 1].text = USECASE_P1
    paras[usecase_idx + 2].text = USECASE_P2
    for idx in range(fig32_idx - 1, usecase_idx + 2, -1):
        if idx > usecase_idx + 2:
            remove_paragraph(doc.paragraphs[idx])

    # Refresh paragraph list after deletions
    paras = doc.paragraphs

    # 3. Section 5.7 key implementation explanation
    key_idx = next(i for i, p in enumerate(paras) if p.text.strip() == "5.7  关键实现说明")
    chapter6_idx = next(i for i, p in enumerate(paras) if p.text.strip() == "第六章  系统测试")
    paras[key_idx + 1].text = KEY_IMPL_P1
    paras[key_idx + 2].text = KEY_IMPL_P2
    for idx in range(chapter6_idx - 1, key_idx + 2, -1):
        if idx > key_idx + 2:
            remove_paragraph(doc.paragraphs[idx])

    doc.save(str(PAPER))
    print(f"updated: {PAPER}")
    print(f"backup: {BACKUP}")


if __name__ == "__main__":
    main()
