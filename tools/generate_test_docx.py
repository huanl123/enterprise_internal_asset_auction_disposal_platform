from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


def zh(text: str) -> str:
    return text.encode("ascii").decode("unicode_escape")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:%s" % edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:%s" % key), str(value))


def set_east_asia_font(run, east_asia="SimSun", latin="Times New Roman"):
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_text(paragraph, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(size)
    run.bold = bold
    return run


def format_cell(cell, text, bold=False, center=True, size=10.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_text(
        paragraph,
        text,
        bold=bold,
        size=size,
        align=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT,
    )
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(
        cell,
        top={"val": "single", "sz": 8, "color": "000000"},
        bottom={"val": "single", "sz": 8, "color": "000000"},
        left={"val": "single", "sz": 8, "color": "000000"},
        right={"val": "single", "sz": 8, "color": "000000"},
    )


def set_document_defaults(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.name = "Times New Roman"
    r_pr = normal._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "SimSun")

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


TABLE_HEADERS = [
    zh("\\u6240\\u5c5e\\u6a21\\u5757"),
    zh("\\u7528\\u4f8b\\u540d\\u79f0"),
    zh("\\u7528\\u4f8b\\u7f16\\u53f7"),
    zh("\\u7528\\u4f8b\\u7b49\\u7ea7"),
    zh("\\u9884\\u7f6e\\u6761\\u4ef6"),
    zh("\\u6d4b\\u8bd5\\u6b65\\u9aa4"),
    zh("\\u9884\\u671f\\u7ed3\\u679c"),
    zh("\\u5b9e\\u9645\\u7ed3\\u679c"),
]


TABLES = [
    (
        zh("\\u88686.2  \\u7528\\u6237\\u4e0e\\u6743\\u9650\\u6a21\\u5757\\u6d4b\\u8bd5"),
        [
            [
                zh("\\u7528\\u6237\\u4e0e\\u6743\\u9650\\u6a21\\u5757"),
                zh("\\u666e\\u901a\\u5458\\u5de5\\u6ce8\\u518c"),
                "User_fun_1",
                "L1",
                zh("\\u7cfb\\u7edf\\u6b63\\u5e38\\u8fd0\\u884c\\uff0c\\u7528\\u6237\\u672a\\u6ce8\\u518c\\u8d26\\u53f7"),
                zh("1\\u3001\\u8f93\\u5165\\u5de5\\u53f7\\u3001\\u5bc6\\u7801\\u3001\\u59d3\\u540d\\u3001\\u8054\\u7cfb\\u65b9\\u5f0f\\u5e76\\u9009\\u62e9\\u90e8\\u95e8\\uff1b2\\u3001\\u70b9\\u51fb\\u201c\\u6ce8\\u518c\\u201d\\u6309\\u94ae"),
                zh("\\u7cfb\\u7edf\\u63d0\\u793a\\u6ce8\\u518c\\u6210\\u529f\\uff0c\\u8d26\\u53f7\\u53ef\\u6b63\\u5e38\\u767b\\u5f55"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u7528\\u6237\\u4e0e\\u6743\\u9650\\u6a21\\u5757"),
                zh("\\u666e\\u901a\\u5458\\u5de5\\u767b\\u5f55"),
                "User_fun_2",
                "L1",
                zh("\\u8d26\\u53f7\\u5df2\\u6ce8\\u518c\\u4e14\\u5904\\u4e8e\\u542f\\u7528\\u72b6\\u6001"),
                zh("1\\u3001\\u8f93\\u5165\\u6b63\\u786e\\u5de5\\u53f7\\u4e0e\\u5bc6\\u7801\\uff1b2\\u3001\\u70b9\\u51fb\\u201c\\u767b\\u5f55\\u201d\\u6309\\u94ae"),
                zh("\\u767b\\u5f55\\u6210\\u529f\\uff0c\\u8df3\\u8f6c\\u81f3\\u7cfb\\u7edf\\u9996\\u9875\\u5e76\\u663e\\u793a\\u5bf9\\u5e94\\u83dc\\u5355"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u7528\\u6237\\u4e0e\\u6743\\u9650\\u6a21\\u5757"),
                zh("\\u7981\\u7528\\u8d26\\u53f7\\u767b\\u5f55"),
                "User_fun_3",
                "L1",
                zh("\\u8d26\\u53f7\\u5df2\\u88ab\\u7cfb\\u7edf\\u7ba1\\u7406\\u5458\\u7981\\u7528"),
                zh("1\\u3001\\u8f93\\u5165\\u6b63\\u786e\\u5de5\\u53f7\\u4e0e\\u5bc6\\u7801\\uff1b2\\u3001\\u70b9\\u51fb\\u201c\\u767b\\u5f55\\u201d\\u6309\\u94ae"),
                zh("\\u7cfb\\u7edf\\u62d2\\u7edd\\u767b\\u5f55\\uff0c\\u63d0\\u793a\\u8d26\\u53f7\\u5df2\\u88ab\\u7981\\u7528"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u7528\\u6237\\u4e0e\\u6743\\u9650\\u6a21\\u5757"),
                zh("\\u89d2\\u8272\\u6743\\u9650\\u9694\\u79bb"),
                "User_fun_4",
                "L1",
                zh("\\u7cfb\\u7edf\\u4e2d\\u5b58\\u5728\\u666e\\u901a\\u5458\\u5de5\\u3001\\u8d44\\u4ea7\\u4e13\\u5458\\u3001\\u8d22\\u52a1\\u4e13\\u5458\\u3001\\u7cfb\\u7edf\\u7ba1\\u7406\\u5458\\u8d26\\u53f7"),
                zh("1\\u3001\\u5206\\u522b\\u4f7f\\u7528\\u4e0d\\u540c\\u89d2\\u8272\\u8d26\\u53f7\\u767b\\u5f55\\uff1b2\\u3001\\u8bbf\\u95ee\\u8d44\\u4ea7\\u3001\\u8d22\\u52a1\\u5ba1\\u6838\\u7b49\\u83dc\\u5355"),
                zh("\\u4e0d\\u540c\\u89d2\\u8272\\u4ec5\\u80fd\\u8bbf\\u95ee\\u4e0e\\u81ea\\u8eab\\u804c\\u8d23\\u5bf9\\u5e94\\u7684\\u83dc\\u5355\\u548c\\u63a5\\u53e3"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u7528\\u6237\\u4e0e\\u6743\\u9650\\u6a21\\u5757"),
                zh("\\u672a\\u767b\\u5f55\\u8bbf\\u95ee\\u53d7\\u9650\\u9875\\u9762"),
                "User_fun_5",
                "L1",
                zh("\\u7528\\u6237\\u672a\\u767b\\u5f55\\u7cfb\\u7edf"),
                zh("1\\u3001\\u76f4\\u63a5\\u8bbf\\u95ee\\u8d44\\u4ea7\\u7ba1\\u7406\\u3001\\u4ed8\\u6b3e\\u5ba1\\u6838\\u7b49\\u4e1a\\u52a1\\u9875\\u9762"),
                zh("\\u7cfb\\u7edf\\u8df3\\u8f6c\\u5230\\u767b\\u5f55\\u9875\\u6216\\u8fd4\\u56de\\u672a\\u6388\\u6743\\u63d0\\u793a"),
                zh("\\u901a\\u8fc7"),
            ],
        ],
    ),
    (
        zh("\\u88686.3  \\u8d44\\u4ea7\\u7ba1\\u7406\\u4e0e\\u5ba1\\u6838\\u6a21\\u5757\\u6d4b\\u8bd5"),
        [
            [
                zh("\\u8d44\\u4ea7\\u7ba1\\u7406\\u4e0e\\u5ba1\\u6838\\u6a21\\u5757"),
                zh("\\u65b0\\u589e\\u8d44\\u4ea7\\u4fe1\\u606f"),
                "Asset_fun_1",
                "L1",
                zh("\\u8d44\\u4ea7\\u4e13\\u5458\\u5df2\\u767b\\u5f55\\u7cfb\\u7edf\\uff0c\\u6298\\u65e7\\u89c4\\u5219\\u5df2\\u5efa\\u7acb"),
                zh("1\\u3001\\u8fdb\\u5165\\u65b0\\u589e\\u8d44\\u4ea7\\u9875\\u9762\\uff1b2\\u3001\\u586b\\u5199\\u8d44\\u4ea7\\u540d\\u79f0\\u3001\\u7f16\\u53f7\\u3001\\u89c4\\u683c\\u3001\\u90e8\\u95e8\\u3001\\u8d2d\\u7f6e\\u65e5\\u671f\\u3001\\u539f\\u503c\\u7b49\\u4fe1\\u606f\\u5e76\\u4e0a\\u4f20\\u56fe\\u7247\\uff1b3\\u3001\\u70b9\\u51fb\\u4fdd\\u5b58"),
                zh("\\u8d44\\u4ea7\\u5efa\\u6863\\u6210\\u529f\\uff0c\\u751f\\u6210\\u8d44\\u4ea7\\u8bb0\\u5f55\\u5e76\\u8fdb\\u5165\\u201c\\u5f85\\u5ba1\\u6838\\u201d\\u72b6\\u6001"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u8d44\\u4ea7\\u7ba1\\u7406\\u4e0e\\u5ba1\\u6838\\u6a21\\u5757"),
                zh("\\u8d44\\u4ea7\\u5fc5\\u586b\\u9879\\u6821\\u9a8c"),
                "Asset_fun_2",
                "L1",
                zh("\\u8d44\\u4ea7\\u4e13\\u5458\\u5df2\\u8fdb\\u5165\\u65b0\\u589e\\u8d44\\u4ea7\\u9875\\u9762"),
                zh("1\\u3001\\u4e0d\\u586b\\u5199\\u5fc5\\u586b\\u9879\\u6216\\u4e0d\\u9009\\u62e9\\u6298\\u65e7\\u89c4\\u5219\\uff1b2\\u3001\\u76f4\\u63a5\\u63d0\\u4ea4"),
                zh("\\u7cfb\\u7edf\\u63d0\\u793a\\u5fc5\\u586b\\u9879\\u4e0d\\u80fd\\u4e3a\\u7a7a\\uff0c\\u963b\\u6b62\\u8d44\\u4ea7\\u4fdd\\u5b58"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u8d44\\u4ea7\\u7ba1\\u7406\\u4e0e\\u5ba1\\u6838\\u6a21\\u5757"),
                zh("\\u8d44\\u4ea7\\u5ba1\\u6838\\u901a\\u8fc7"),
                "Asset_fun_3",
                "L1",
                zh("\\u5b58\\u5728\\u72b6\\u6001\\u4e3a\\u201c\\u5f85\\u5ba1\\u6838\\u201d\\u7684\\u8d44\\u4ea7\\uff0c\\u8d22\\u52a1\\u4e13\\u5458\\u5df2\\u767b\\u5f55"),
                zh("1\\u3001\\u5728\\u5ba1\\u6838\\u9875\\u9762\\u8bbe\\u7f6e\\u8d77\\u62cd\\u4ef7\\u3001\\u4fdd\\u7559\\u4ef7\\uff1b2\\u3001\\u586b\\u5199\\u5ba1\\u6838\\u610f\\u89c1\\uff1b3\\u3001\\u70b9\\u51fb\\u201c\\u901a\\u8fc7\\u201d"),
                zh("\\u8d44\\u4ea7\\u72b6\\u6001\\u66f4\\u65b0\\u4e3a\\u201c\\u5f85\\u62cd\\u5356\\u201d\\uff0c\\u751f\\u6210\\u5ba1\\u6838\\u5386\\u53f2\\u8bb0\\u5f55"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u8d44\\u4ea7\\u7ba1\\u7406\\u4e0e\\u5ba1\\u6838\\u6a21\\u5757"),
                zh("\\u8d44\\u4ea7\\u5ba1\\u6838\\u62d2\\u7edd"),
                "Asset_fun_4",
                "L1",
                zh("\\u5b58\\u5728\\u72b6\\u6001\\u4e3a\\u201c\\u5f85\\u5ba1\\u6838\\u201d\\u7684\\u8d44\\u4ea7\\uff0c\\u8d22\\u52a1\\u4e13\\u5458\\u5df2\\u767b\\u5f55"),
                zh("1\\u3001\\u5728\\u5ba1\\u6838\\u9875\\u9762\\u586b\\u5199\\u62d2\\u7edd\\u539f\\u56e0\\uff1b2\\u3001\\u70b9\\u51fb\\u201c\\u62d2\\u7edd\\u201d"),
                zh("\\u7cfb\\u7edf\\u8bb0\\u5f55\\u62d2\\u7edd\\u539f\\u56e0\\uff0c\\u8d44\\u4ea7\\u9000\\u56de\\u4fee\\u6539\\u6d41\\u7a0b"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u8d44\\u4ea7\\u7ba1\\u7406\\u4e0e\\u5ba1\\u6838\\u6a21\\u5757"),
                zh("\\u91cd\\u65b0\\u8ba1\\u7b97\\u8d44\\u4ea7\\u4ef7\\u503c"),
                "Asset_fun_5",
                "L1",
                zh("\\u8d44\\u4ea7\\u672a\\u53c2\\u4e0e\\u62cd\\u5356\\uff0c\\u5141\\u8bb8\\u91cd\\u65b0\\u53d1\\u8d77\\u5ba1\\u6838"),
                zh("1\\u3001\\u70b9\\u51fb\\u201c\\u91cd\\u65b0\\u8ba1\\u7b97\\u201d\\u6309\\u94ae\\uff1b2\\u3001\\u91cd\\u65b0\\u89e6\\u53d1\\u5b9a\\u4ef7\\u8bc4\\u4f30"),
                zh("\\u7cfb\\u7edf\\u91cd\\u65b0\\u8ba1\\u7b97\\u5f53\\u524d\\u4ef7\\u503c\\u4e0e\\u5efa\\u8bae\\u8d77\\u62cd\\u4ef7\\uff0c\\u5e76\\u91cd\\u65b0\\u8fdb\\u5165\\u5ba1\\u6838\\u6d41\\u7a0b"),
                zh("\\u901a\\u8fc7"),
            ],
        ],
    ),
    (
        zh("\\u88686.4  \\u62cd\\u5356\\u4e0e\\u4ea4\\u6613\\u6a21\\u5757\\u6d4b\\u8bd5"),
        [
            [
                zh("\\u62cd\\u5356\\u4e0e\\u4ea4\\u6613\\u6a21\\u5757"),
                zh("\\u521b\\u5efa\\u62cd\\u5356\\u6d3b\\u52a8"),
                "Auction_fun_1",
                "L1",
                zh("\\u5b58\\u5728\\u72b6\\u6001\\u4e3a\\u201c\\u5f85\\u62cd\\u5356\\u201d\\u7684\\u8d44\\u4ea7\\uff0c\\u8d44\\u4ea7\\u4e13\\u5458\\u5df2\\u767b\\u5f55"),
                zh("1\\u3001\\u586b\\u5199\\u62cd\\u5356\\u6807\\u9898\\u3001\\u8d77\\u62cd\\u4ef7\\u3001\\u52a0\\u4ef7\\u5e45\\u5ea6\\u3001\\u5f00\\u59cb\\u65f6\\u95f4\\u3001\\u7ed3\\u675f\\u65f6\\u95f4\\uff1b2\\u3001\\u70b9\\u51fb\\u53d1\\u5e03"),
                zh("\\u62cd\\u5356\\u6d3b\\u52a8\\u521b\\u5efa\\u6210\\u529f\\uff0c\\u8d44\\u4ea7\\u8fdb\\u5165\\u62cd\\u5356\\u6d41\\u7a0b"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u62cd\\u5356\\u4e0e\\u4ea4\\u6613\\u6a21\\u5757"),
                zh("\\u5458\\u5de5\\u6b63\\u5e38\\u51fa\\u4ef7"),
                "Auction_fun_2",
                "L1",
                zh("\\u62cd\\u5356\\u6d3b\\u52a8\\u5904\\u4e8e\\u8fdb\\u884c\\u4e2d\\uff0c\\u666e\\u901a\\u5458\\u5de5\\u5177\\u5907\\u7ade\\u62cd\\u8d44\\u683c"),
                zh("1\\u3001\\u5728\\u62cd\\u5356\\u8be6\\u60c5\\u9875\\u9762\\u8f93\\u5165\\u5408\\u6cd5\\u91d1\\u989d\\uff1b2\\u3001\\u70b9\\u51fb\\u63d0\\u4ea4\\u51fa\\u4ef7"),
                zh("\\u7cfb\\u7edf\\u63a5\\u53d7\\u51fa\\u4ef7\\uff0c\\u66f4\\u65b0\\u5f53\\u524d\\u6700\\u9ad8\\u4ef7\\u4e0e\\u7ade\\u4ef7\\u8bb0\\u5f55"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u62cd\\u5356\\u4e0e\\u4ea4\\u6613\\u6a21\\u5757"),
                zh("\\u975e\\u6cd5\\u51fa\\u4ef7\\u6821\\u9a8c"),
                "Auction_fun_3",
                "L1",
                zh("\\u62cd\\u5356\\u6d3b\\u52a8\\u5904\\u4e8e\\u8fdb\\u884c\\u4e2d"),
                zh("1\\u3001\\u8f93\\u5165\\u4f4e\\u4e8e\\u5f53\\u524d\\u4ef7\\u52a0\\u6700\\u5c0f\\u52a0\\u4ef7\\u5e45\\u5ea6\\u7684\\u91d1\\u989d\\uff1b2\\u3001\\u70b9\\u51fb\\u63d0\\u4ea4"),
                zh("\\u7cfb\\u7edf\\u62d2\\u7edd\\u51fa\\u4ef7\\uff0c\\u63d0\\u793a\\u91d1\\u989d\\u4e0d\\u7b26\\u5408\\u89c4\\u5219"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u62cd\\u5356\\u4e0e\\u4ea4\\u6613\\u6a21\\u5757"),
                zh("\\u62cd\\u5356\\u7ed3\\u675f\\u751f\\u6210\\u4ea4\\u6613"),
                "Auction_fun_4",
                "L1",
                zh("\\u62cd\\u5356\\u5230\\u8fbe\\u7ed3\\u675f\\u65f6\\u95f4\\u4e14\\u5b58\\u5728\\u6709\\u6548\\u6700\\u9ad8\\u4ef7"),
                zh("1\\u3001\\u7531\\u7cfb\\u7edf\\u7ed3\\u675f\\u62cd\\u5356\\uff1b2\\u3001\\u81ea\\u52a8\\u5224\\u5b9a\\u7ed3\\u679c"),
                zh("\\u751f\\u6210\\u5f85\\u786e\\u8ba4\\u4ea4\\u6613\\u5355\\uff0c\\u4e2d\\u6807\\u5458\\u5de5\\u53ef\\u5728\\u89c4\\u5b9a\\u65f6\\u95f4\\u5185\\u786e\\u8ba4\\u6210\\u4ea4"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u62cd\\u5356\\u4e0e\\u4ea4\\u6613\\u6a21\\u5757"),
                zh("\\u6d41\\u62cd\\u5904\\u7406"),
                "Auction_fun_5",
                "L1",
                zh("\\u62cd\\u5356\\u7ed3\\u675f\\u4e14\\u65e0\\u4eba\\u51fa\\u4ef7\\u6216\\u6700\\u9ad8\\u4ef7\\u672a\\u8fbe\\u4fdd\\u7559\\u4ef7"),
                zh("1\\u3001\\u7531\\u7cfb\\u7edf\\u6267\\u884c\\u62cd\\u5356\\u7ed3\\u7b97\\u903b\\u8f91"),
                zh("\\u62cd\\u5356\\u7ed3\\u679c\\u8bb0\\u4e3a\\u6d41\\u62cd\\uff0c\\u8d44\\u4ea7\\u72b6\\u6001\\u6062\\u590d\\u4e3a\\u201c\\u5f85\\u62cd\\u5356\\u201d"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u62cd\\u5356\\u4e0e\\u4ea4\\u6613\\u6a21\\u5757"),
                zh("\\u4e2d\\u6807\\u5458\\u5de5\\u786e\\u8ba4\\u6210\\u4ea4"),
                "Auction_fun_6",
                "L1",
                zh("\\u5b58\\u5728\\u5f85\\u786e\\u8ba4\\u4ea4\\u6613\\u5355\\uff0c\\u4e2d\\u6807\\u5458\\u5de5\\u5df2\\u767b\\u5f55"),
                zh("1\\u3001\\u5728\\u201c\\u6211\\u7684\\u7ade\\u62cd\\u201d\\u9875\\u9762\\u70b9\\u51fb\\u201c\\u786e\\u8ba4\\u6210\\u4ea4\\u201d"),
                zh("\\u4ea4\\u6613\\u8fdb\\u5165\\u4ed8\\u6b3e\\u5f85\\u5ba1\\u6838\\u6d41\\u7a0b\\uff0c\\u751f\\u6210\\u5bf9\\u5e94\\u8bb0\\u5f55"),
                zh("\\u901a\\u8fc7"),
            ],
        ],
    ),
    (
        zh("\\u88686.5  \\u4ed8\\u6b3e\\u5ba1\\u6838\\u4e0e\\u5904\\u7f6e\\u5f52\\u6863\\u6a21\\u5757\\u6d4b\\u8bd5"),
        [
            [
                zh("\\u4ed8\\u6b3e\\u5ba1\\u6838\\u4e0e\\u5904\\u7f6e\\u5f52\\u6863\\u6a21\\u5757"),
                zh("\\u8d22\\u52a1\\u786e\\u8ba4\\u6536\\u6b3e\\u901a\\u8fc7"),
                "Finance_fun_1",
                "L1",
                zh("\\u5b58\\u5728\\u4ed8\\u6b3e\\u5f85\\u5ba1\\u6838\\u4ea4\\u6613\\u5355\\uff0c\\u8d22\\u52a1\\u4e13\\u5458\\u5df2\\u767b\\u5f55"),
                zh("1\\u3001\\u5728\\u4ed8\\u6b3e\\u5f85\\u5ba1\\u6838\\u9875\\u9762\\u4e0a\\u4f20\\u4ed8\\u6b3e\\u51ed\\u8bc1\\uff1b2\\u3001\\u586b\\u5199\\u5ba1\\u6838\\u5907\\u6ce8\\uff1b3\\u3001\\u70b9\\u51fb\\u201c\\u786e\\u8ba4\\u6536\\u6b3e\\u201d"),
                zh("\\u4ea4\\u6613\\u4ed8\\u6b3e\\u72b6\\u6001\\u66f4\\u65b0\\u4e3a approved\\uff0c\\u8d44\\u4ea7\\u72b6\\u6001\\u6d41\\u8f6c\\u81f3\\u201c\\u5f85\\u5904\\u7f6e\\u201d"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u4ed8\\u6b3e\\u5ba1\\u6838\\u4e0e\\u5904\\u7f6e\\u5f52\\u6863\\u6a21\\u5757"),
                zh("\\u4ed8\\u6b3e\\u51ed\\u8bc1\\u4e3a\\u7a7a\\u6821\\u9a8c"),
                "Finance_fun_2",
                "L1",
                zh("\\u5b58\\u5728\\u4ed8\\u6b3e\\u5f85\\u5ba1\\u6838\\u4ea4\\u6613\\u5355"),
                zh("1\\u3001\\u4e0d\\u4e0a\\u4f20\\u4ed8\\u6b3e\\u51ed\\u8bc1\\uff1b2\\u3001\\u76f4\\u63a5\\u63d0\\u4ea4\\u786e\\u8ba4\\u6536\\u6b3e"),
                zh("\\u7cfb\\u7edf\\u63d0\\u793a\\u4ed8\\u6b3e\\u51ed\\u8bc1\\u4e0d\\u80fd\\u4e3a\\u7a7a\\uff0c\\u963b\\u6b62\\u63d0\\u4ea4"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u4ed8\\u6b3e\\u5ba1\\u6838\\u4e0e\\u5904\\u7f6e\\u5f52\\u6863\\u6a21\\u5757"),
                zh("\\u8d22\\u52a1\\u5ba1\\u6838\\u4e0d\\u901a\\u8fc7"),
                "Finance_fun_3",
                "L1",
                zh("\\u5b58\\u5728\\u4ed8\\u6b3e\\u5f85\\u5ba1\\u6838\\u4ea4\\u6613\\u5355\\uff0c\\u8d22\\u52a1\\u4e13\\u5458\\u5df2\\u767b\\u5f55"),
                zh("1\\u3001\\u586b\\u5199\\u62d2\\u7edd\\u539f\\u56e0\\uff1b2\\u3001\\u70b9\\u51fb\\u201c\\u4e0d\\u901a\\u8fc7\\u201d"),
                zh("\\u7cfb\\u7edf\\u53d6\\u6d88\\u672c\\u6b21\\u4ea4\\u6613\\uff0c\\u8d44\\u4ea7\\u6062\\u590d\\u4e3a\\u201c\\u5f85\\u62cd\\u5356\\u201d\\uff0c\\u4e2d\\u6807\\u8005\\u53d7\\u5230\\u7ade\\u62cd\\u9650\\u5236"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u4ed8\\u6b3e\\u5ba1\\u6838\\u4e0e\\u5904\\u7f6e\\u5f52\\u6863\\u6a21\\u5757"),
                zh("\\u786e\\u8ba4\\u8d44\\u4ea7\\u5904\\u7f6e"),
                "Disposal_fun_1",
                "L1",
                zh("\\u5b58\\u5728\\u72b6\\u6001\\u4e3a\\u201c\\u5f85\\u5904\\u7f6e\\u201d\\u7684\\u8d44\\u4ea7\\uff0c\\u8d44\\u4ea7\\u4e13\\u5458\\u5df2\\u767b\\u5f55"),
                zh("1\\u3001\\u5728\\u5904\\u7f6e\\u7ba1\\u7406\\u9875\\u9762\\u4e0a\\u4f20\\u5904\\u7f6e\\u51ed\\u8bc1\\uff1b2\\u3001\\u586b\\u5199\\u5904\\u7f6e\\u5907\\u6ce8\\uff1b3\\u3001\\u70b9\\u51fb\\u201c\\u786e\\u8ba4\\u5904\\u7f6e\\u201d"),
                zh("\\u8d44\\u4ea7\\u72b6\\u6001\\u66f4\\u65b0\\u4e3a\\u201c\\u5df2\\u5904\\u7f6e\\u201d\\uff0c\\u5904\\u7f6e\\u8bb0\\u5f55\\u4fdd\\u5b58\\u6210\\u529f"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u4ed8\\u6b3e\\u5ba1\\u6838\\u4e0e\\u5904\\u7f6e\\u5f52\\u6863\\u6a21\\u5757"),
                zh("\\u751f\\u6210\\u5e76\\u67e5\\u8be2\\u5904\\u7f6e\\u6863\\u6848"),
                "Disposal_fun_2",
                "L1",
                zh("\\u8d44\\u4ea7\\u5df2\\u5b8c\\u6210\\u5904\\u7f6e"),
                zh("1\\u3001\\u5728\\u8d44\\u4ea7\\u5904\\u7f6e\\u6863\\u6848\\u67e5\\u8be2\\u9875\\u9762\\u6309\\u8d44\\u4ea7\\u7f16\\u53f7\\u6216\\u540d\\u79f0\\u68c0\\u7d22"),
                zh("\\u7cfb\\u7edf\\u663e\\u793a\\u5f52\\u6863\\u65f6\\u95f4\\u3001\\u72b6\\u6001\\u53ca\\u5b8c\\u6574\\u6863\\u6848\\u4fe1\\u606f"),
                zh("\\u901a\\u8fc7"),
            ],
        ],
    ),
    (
        zh("\\u88686.6  \\u8d44\\u4ea7\\u5904\\u7f6e\\u7edf\\u8ba1\\u6a21\\u5757\\u6d4b\\u8bd5"),
        [
            [
                zh("\\u8d44\\u4ea7\\u5904\\u7f6e\\u7edf\\u8ba1\\u6a21\\u5757"),
                zh("\\u6309\\u6708\\u7edf\\u8ba1\\u67e5\\u8be2"),
                "Stat_fun_1",
                "L1",
                zh("\\u7cfb\\u7edf\\u4e2d\\u5df2\\u5b58\\u5728\\u8d44\\u4ea7\\u5904\\u7f6e\\u6570\\u636e"),
                zh("1\\u3001\\u5728\\u8d44\\u4ea7\\u5904\\u7f6e\\u7edf\\u8ba1\\u9875\\u9762\\u9009\\u62e9\\u201c\\u6309\\u6708\\u201d\\uff1b2\\u3001\\u6307\\u5b9a\\u65f6\\u95f4\\u70b9\\u540e\\u70b9\\u51fb\\u201c\\u5237\\u65b0\\u201d"),
                zh("\\u9875\\u9762\\u6b63\\u786e\\u5c55\\u793a\\u62cd\\u5356\\u7ed3\\u675f\\u6570\\u3001\\u6210\\u4ea4\\u91d1\\u989d\\u3001\\u6d41\\u62cd\\u7387\\u53ca\\u8d8b\\u52bf\\u56fe\\u6570\\u636e"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u8d44\\u4ea7\\u5904\\u7f6e\\u7edf\\u8ba1\\u6a21\\u5757"),
                zh("\\u6309\\u5b63\\u7edf\\u8ba1\\u67e5\\u8be2"),
                "Stat_fun_2",
                "L1",
                zh("\\u7cfb\\u7edf\\u4e2d\\u5df2\\u5b58\\u5728\\u8d44\\u4ea7\\u5904\\u7f6e\\u6570\\u636e"),
                zh("1\\u3001\\u5728\\u8d44\\u4ea7\\u5904\\u7f6e\\u7edf\\u8ba1\\u9875\\u9762\\u9009\\u62e9\\u201c\\u6309\\u5b63\\u201d\\uff1b2\\u3001\\u70b9\\u51fb\\u201c\\u5237\\u65b0\\u201d"),
                zh("\\u7cfb\\u7edf\\u6309\\u5b63\\u5ea6\\u6c47\\u603b\\u5904\\u7f6e\\u6570\\u91cf\\u3001\\u6210\\u4ea4\\u91d1\\u989d\\u548c\\u6d41\\u62cd\\u7387"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u8d44\\u4ea7\\u5904\\u7f6e\\u7edf\\u8ba1\\u6a21\\u5757"),
                zh("\\u6309\\u5e74\\u7edf\\u8ba1\\u67e5\\u8be2"),
                "Stat_fun_3",
                "L1",
                zh("\\u7cfb\\u7edf\\u4e2d\\u5df2\\u5b58\\u5728\\u8d44\\u4ea7\\u5904\\u7f6e\\u6570\\u636e"),
                zh("1\\u3001\\u5728\\u8d44\\u4ea7\\u5904\\u7f6e\\u7edf\\u8ba1\\u9875\\u9762\\u9009\\u62e9\\u201c\\u6309\\u5e74\\u201d\\uff1b2\\u3001\\u70b9\\u51fb\\u201c\\u5237\\u65b0\\u201d"),
                zh("\\u7cfb\\u7edf\\u6309\\u5e74\\u5ea6\\u6c47\\u603b\\u7edf\\u8ba1\\u7ed3\\u679c\\u5e76\\u66f4\\u65b0\\u8d8b\\u52bf\\u56fe"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u8d44\\u4ea7\\u5904\\u7f6e\\u7edf\\u8ba1\\u6a21\\u5757"),
                zh("\\u90e8\\u95e8\\u5904\\u7f6e\\u6c47\\u603b\\u663e\\u793a"),
                "Stat_fun_4",
                "L1",
                zh("\\u5f53\\u524d\\u7edf\\u8ba1\\u5468\\u671f\\u5185\\u5b58\\u5728\\u591a\\u4e2a\\u90e8\\u95e8\\u5904\\u7f6e\\u8bb0\\u5f55"),
                zh("1\\u3001\\u67e5\\u770b\\u9875\\u9762\\u4e0b\\u65b9\\u90e8\\u95e8\\u5904\\u7f6e\\u6c47\\u603b\\u8868"),
                zh("\\u7cfb\\u7edf\\u663e\\u793a\\u5404\\u90e8\\u95e8\\u5904\\u7f6e\\u6570\\u91cf\\u3001\\u6210\\u4ea4\\u91d1\\u989d\\u53ca\\u6d41\\u62cd\\u7387"),
                zh("\\u901a\\u8fc7"),
            ],
            [
                zh("\\u8d44\\u4ea7\\u5904\\u7f6e\\u7edf\\u8ba1\\u6a21\\u5757"),
                zh("\\u5bfc\\u51fa\\u7edf\\u8ba1\\u6570\\u636e"),
                "Stat_fun_5",
                "L1",
                zh("\\u5f53\\u524d\\u7edf\\u8ba1\\u9875\\u9762\\u5df2\\u751f\\u6210\\u6709\\u6548\\u7edf\\u8ba1\\u7ed3\\u679c"),
                zh("1\\u3001\\u70b9\\u51fb\\u201c\\u5bfc\\u51fa\\u201d\\u6309\\u94ae"),
                zh("\\u7cfb\\u7edf\\u6210\\u529f\\u5bfc\\u51fa\\u7edf\\u8ba1\\u7ed3\\u679c\\u6587\\u4ef6"),
                zh("\\u901a\\u8fc7"),
            ],
        ],
    ),
]


def set_column_widths(table):
    widths = [Cm(1.2), Cm(2.1), Cm(2.2), Cm(1.5), Cm(3.3), Cm(4.0), Cm(3.3), Cm(1.5)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def create_table(doc, title, rows):
    title_p = doc.add_paragraph()
    set_paragraph_text(title_p, title, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    table = doc.add_table(rows=1, cols=len(TABLE_HEADERS))
    table.style = "Table Grid"
    table.autofit = False
    set_column_widths(table)

    header_cells = table.rows[0].cells
    for idx, header in enumerate(TABLE_HEADERS):
        format_cell(header_cells[idx], header, bold=False, center=True, size=10.5)

    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            format_cell(row.cells[idx], value, bold=False, center=(idx in (0, 2, 3, 7)), size=10.5)

    doc.add_paragraph("")


def main():
    output = Path("images") / "test.docx"
    doc = Document()
    set_document_defaults(doc)

    title = doc.add_paragraph()
    set_paragraph_text(
        title,
        zh("\\u6d4b\\u8bd5\\u7528\\u4f8b"),
        bold=True,
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph("")

    for table_title, rows in TABLES:
        create_table(doc, table_title, rows)

    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
