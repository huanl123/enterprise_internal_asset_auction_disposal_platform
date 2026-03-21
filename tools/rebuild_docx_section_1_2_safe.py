import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SRC = Path(r"E:\Study\code\Java\enterprise_internal_asset_auction_disposal_platform\images\2215304659 庄采润 论文.bak_before_rewrite_1_2.docx")
TARGET = Path(r"E:\Study\code\Java\enterprise_internal_asset_auction_disposal_platform\images\2215304659 庄采润 论文.docx")
BROKEN_BACKUP = Path(r"E:\Study\code\Java\enterprise_internal_asset_auction_disposal_platform\images\2215304659 庄采润 论文.bak_corrupt_before_safe_fix.docx")


NEW_PARAGRAPHS = [
    [
        ("国内关于企业资产管理与相关信息系统的研究，已从单纯台账管理逐步转向流程化、信息化和协同化建设。前后端分离和 B/S 架构方面的研究表明，模块解耦、接口统一与页面交互优化能够有效提升系统的扩展性和用户体验，为中后台业务系统建设提供了成熟思路", False),
        ("[2,4,8]", True),
        ("。在此基础上，张文、李明等围绕固定资产信息管理系统展开实践，已实现资产信息录入、折旧计算、审批提醒与基础报表等功能", False),
        ("[3,14]", True),
        ("。不过，这类研究大多聚焦资产登记和日常管理，对内部拍卖、成交确认、付款审核与处置归档等后续环节覆盖不足，尚未形成面向处置业务的完整闭环。", False),
    ],
    [
        ("从管理机制和业务需求层面看，罗洁、周悦、王晗、沈菲、杨静及张伟等研究分别从固定资产管理优化、全生命周期成本采集、内控信息化、内部控制改进以及数字化技术应用等角度进行了分析", False),
        ("[6,9,11-13,15]", True),
        ("。相关成果普遍认为，资产管理系统应加强数据贯通、过程留痕和制度约束，提升审批透明度和管理效率。葛辉关于校园闲置物品交易平台的研究，则从交易流程与用户交互角度提供了参考", False),
        ("[1]", True),
        ("。但综合来看，国内现有研究虽已在资产录入、价值计算、基础交易和信息统计方面积累了一定成果，却仍存在拍卖流程自动化不足、保留价与异常处理机制不完善、角色权限粒度不够细以及多部门协同支持不足等问题。", False),
    ],
    [
        ("国外在相关领域起步较早，研究更强调资产全生命周期治理、数据驱动决策以及流程自动化。CHEN 和 LUO 基于 Spring Boot 的物业管理系统研究表明，模块化架构和实时数据更新机制能够有效提升业务处理效率与用户体验", False),
        ("[5]", True),
        ("；Pakkala 等从工业生产资产信息管理出发，强调自动化采集与协同机制对提升资产信息质量和使用效率的重要作用", False),
        ("[7]", True),
        ("；Crespo Marquez 等则从更宏观的角度总结了资产管理技术、评估方法与行业应用的发展趋势，指出信息系统集成、资产评估和流程协同是资产治理的重要方向", False),
        ("[10]", True),
        ("。这些研究为本课题提供了方法借鉴，但已有成果多服务于物业、工业设备或广义资产管理场景，直接面向企业内部废旧资产拍卖与处置全过程的一体化平台研究仍相对有限。", False),
    ],
    [
        ("从具体实现技术来看，MVC 与三层架构、Java Web 技术体系、Vue 前端框架以及 Spring Boot 开发模式已经较为成熟，能够为多角色协同业务系统提供稳定的工程基础", False),
        ("[16-20]", True),
        ("。结合开题报告和任务书中的研究目标可以看出，企业废旧资产内部拍卖与处置平台不仅需要具备一般资产管理系统的数据维护能力，还需要进一步解决自动定价与保留价控制、拍卖结果判定、超时违约处理、付款凭证审核、处置归档锁定及统计分析等专门问题。因此，面向企业内部场景构建一套兼顾流程闭环、权限控制、业务留痕和数据分析能力的平台，仍具有较强的研究价值与实践意义。", False),
    ],
]


def set_east_asia_font(run, font_name: str) -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def copy_paragraph_format(dst, src) -> None:
    dst.style = src.style
    dst.paragraph_format.left_indent = src.paragraph_format.left_indent
    dst.paragraph_format.right_indent = src.paragraph_format.right_indent
    dst.paragraph_format.first_line_indent = src.paragraph_format.first_line_indent
    dst.paragraph_format.space_before = src.paragraph_format.space_before
    dst.paragraph_format.space_after = src.paragraph_format.space_after
    dst.paragraph_format.line_spacing = src.paragraph_format.line_spacing
    dst.paragraph_format.line_spacing_rule = src.paragraph_format.line_spacing_rule
    dst.paragraph_format.alignment = src.paragraph_format.alignment


def remove_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)


def main() -> None:
    if TARGET.exists() and not BROKEN_BACKUP.exists():
        shutil.copy2(TARGET, BROKEN_BACKUP)

    doc = Document(str(SRC))
    body_start = next(i for i, p in enumerate(doc.paragraphs) if i > 100 and "第一章" in p.text and "绪论" in p.text)
    h12 = next(i for i, p in enumerate(doc.paragraphs[body_start:], body_start) if p.text.strip() == "1.2  国内外研究现状")
    h13 = next(i for i, p in enumerate(doc.paragraphs[h12 + 1:], h12 + 1) if p.text.strip() == "1.3  研究目标与主要内容")

    template = doc.paragraphs[h12 + 1]
    heading_13 = doc.paragraphs[h13]

    to_remove = list(doc.paragraphs[h12 + 1:h13])
    for para in to_remove:
        remove_paragraph(para)

    base_font_name = None
    base_font_size = None
    for run in template.runs:
        if run.text:
            base_font_name = run.font.name or "宋体"
            base_font_size = run.font.size
            break
    if base_font_name is None:
        base_font_name = "宋体"

    for segments in NEW_PARAGRAPHS:
        para = heading_13.insert_paragraph_before()
        copy_paragraph_format(para, template)
        for text, superscript in segments:
            run = para.add_run(text)
            run.font.name = "Times New Roman" if superscript else base_font_name
            set_east_asia_font(run, "Times New Roman" if superscript else base_font_name)
            if base_font_size is not None:
                run.font.size = base_font_size
            run.font.superscript = superscript

    doc.save(str(TARGET))
    print(f"rebuilt from backup: {SRC}")
    print(f"saved target: {TARGET}")
    print(f"backup current broken copy: {BROKEN_BACKUP}")


if __name__ == "__main__":
    main()
