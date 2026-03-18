from __future__ import annotations

import argparse

from docx import Document


NEW_REF_20 = (
    "[20] 陈佳, 罗旭莲. 基于Spring Boot的智慧物业管理系统设计与实现[J]. "
    "Academic Journal of Science and Technology, 2022."
)


def insert_paragraph_after(paragraph, text: str) -> None:
    new_p = paragraph._p.addnext(paragraph._p.__class__())
    # The above creates an empty element; easiest is to let python-docx re-wrap:
    # but we can simply access it through a temporary Paragraph wrapper by adding to document body.
    # Instead, use the paragraph's parent to create a new paragraph and move it.
    raise NotImplementedError


def _insert_para_after(paragraph, text: str) -> None:
    # Minimal XML insertion that preserves style reasonably.
    new_p = paragraph._p.addnext(paragraph._p.__class__())
    # Wrap via python-docx's Paragraph class is non-trivial; easiest is to set text through a
    # temporary Document paragraph then swap nodes. We'll do direct runs by assigning .text later.
    # However, python-docx doesn't expose paragraph wrapper for arbitrary oxml. We'll just set
    # text by creating a new paragraph at end and moving it.
    doc_part = paragraph.part
    doc = Document(doc_part.package.part_related_by(r_id=None))  # type: ignore[arg-type]
    raise NotImplementedError


def patch_doc(path: str) -> None:
    doc = Document(path)

    # 1) Update the Chapter 2.6 citation range [13-20] -> [13-19,21]
    for p in doc.paragraphs:
        if "系统安全与接口规范参考" in (p.text or "") and "[13-20]" in (p.text or ""):
            for r in p.runs:
                if "[13-20]" in r.text:
                    r.text = r.text.replace("[13-20]", "[13-19,21]")
                    break
            break

    # 2) Update the OWASP sentence in security testing section to cite [21] and add [20]
    for p in doc.paragraphs:
        if "接口级安全风险评估" in (p.text or "") and "OWASP API Security Top 10" in (p.text or ""):
            # Update the lead-in run to include the new [20] citation and connect to OWASP.
            for r in p.runs:
                if "同时，接口级安全风险评估可参考" in r.text:
                    r.text = (
                        "同时，接口级安全风险评估可参考基于Spring Boot的智慧物业管理系统设计与实现等相关研究中"
                        "的安全与权限控制思路[20]，并结合"
                    )
                    break
            # Update the trailing citation run [20] -> [21]
            for r in p.runs:
                if r.text.strip() == "[20]":
                    r.text = "[21]"
                    break
            break

    # 3) Update reference [20] content and move old [20] (OWASP) to [21]
    # Locate the last "参考文献" heading, then search for paragraphs starting with "[20]".
    ref_start = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t.startswith("参") and ("文" in t) and ("献" in t):
            ref_start = i
    if ref_start is not None:
        refs = doc.paragraphs[ref_start + 1 :]
        p20 = None
        existing21 = False
        for p in refs:
            t = (p.text or "").strip()
            if t.startswith("[21]"):
                existing21 = True
            if t.startswith("[20]"):
                p20 = p
                break

        if p20 is not None:
            old20 = (p20.text or "").strip()
            # Replace [20] line
            p20.text = NEW_REF_20
            # Insert [21] line if missing
            if (not existing21) and old20.startswith("[20]"):
                new21 = "[21]" + old20[len("[20]") :]
                # Insert a new paragraph right after p20 (before appendix heading).
                p20.insert_paragraph_after(new21)

    doc.save(path)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", nargs="+", help="One or more thesis .docx files to patch in-place.")
    args = parser.parse_args()
    for p in args.docx:
        patch_doc(p)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

