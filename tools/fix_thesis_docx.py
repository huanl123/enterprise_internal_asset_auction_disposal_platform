from __future__ import annotations

import argparse
from copy import deepcopy
from dataclasses import dataclass

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@dataclass(frozen=True)
class HeadingFix:
    old_no: str
    old_title: str
    new_no: str
    new_title: str


CH5_FIXES: tuple[HeadingFix, ...] = (
    HeadingFix("5.16", "本章扩展小结", "5.5", "经济与管理效益分析"),
    HeadingFix("5.15", "经济与管理效益分析", "5.6", "系统部署与运维可行性"),
    HeadingFix("5.14", "系统部署与运维可行性", "5.7", "数据统计能力评估"),
    HeadingFix("5.13", "数据统计能力评估", "5.8", "典型测试用例分析"),
    HeadingFix("5.12", "典型测试用例分析", "5.9", "核心业务规则细化说明"),
    HeadingFix("5.11", "核心业务规则细化说明", "5.10", "扩展研究价值说明"),
    HeadingFix("5.20", "扩展研究价值说明", "5.11", "用户体验与交互优化"),
    HeadingFix("5.19", "用户体验与交互优化", "5.12", "缺陷管理与质量改进"),
    HeadingFix("5.18", "缺陷管理与质量改进", "5.13", "软件工程过程管理实践"),
    HeadingFix("5.17", "软件工程过程管理实践", "5.14", "项目风险与应对策略补充"),
    HeadingFix("5.21", "项目风险与应对策略补充", "5.15", "本章小结"),
)

APPENDIX_ENDPOINT_REPLACEMENTS: dict[str, str] = {
    "（3）POST /api/auctions：创建拍卖活动。": "（3）POST /api/auction：创建拍卖活动。",
    "（4）POST /api/auctions/{id}/bids：提交竞价。": "（4）POST /api/auction/{id}/bid：提交竞价。",
    "（5）POST /api/transactions/{id}/confirm：中标者确认成交。": "（5）POST /api/auction/{id}/confirm：中标者确认成交。",
    "（6）POST /api/finance/payments/{id}/review：财务审核付款凭证。": "（6）POST /api/finance/transactions/{id}/payment：财务审核付款凭证。",
    "（7）POST /api/disposals/{id}/complete：资产专员确认处置归档。": "（7）POST /api/disposal/assets/{id}/confirm：资产专员确认处置归档。",
}


def _set_outline_level(paragraph, level: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        pPr.append(outline)
    outline.set(qn("w:val"), str(level))


def _try_fix_body_heading(paragraph, fix: HeadingFix) -> bool:
    text = (paragraph.text or "").strip()
    if text != f"{fix.old_no}  {fix.old_title}":
        return False
    if len(paragraph.runs) < 2:
        return False
    paragraph.runs[0].text = f"{fix.new_no}  "
    paragraph.runs[1].text = fix.new_title
    return True


def _try_fix_toc_heading(paragraph, fix: HeadingFix) -> bool:
    """
    TOC entries are often hyperlink-based; paragraph.runs may be empty.
    We update underlying w:t nodes to preserve hyperlinks/tab leaders.
    """
    text = (paragraph.text or "").strip()
    # TOC text ends with a tab + page number; match by prefix only.
    if not text.startswith(f"{fix.old_no}  {fix.old_title}\t"):
        return False

    texts = paragraph._p.xpath(".//w:t")
    if len(texts) < 2:
        return False
    texts[0].text = f"{fix.new_no}  "
    texts[1].text = fix.new_title
    return True


def fix_docx(input_path: str, output_path: str) -> None:
    doc = Document(input_path)

    # 1) Fix Chapter 5 headings (both TOC entries and body headings).
    toc_fixed = 0
    body_fixed = 0
    for p in doc.paragraphs:
        for fix in CH5_FIXES:
            if _try_fix_toc_heading(p, fix):
                toc_fixed += 1
                break
            if _try_fix_body_heading(p, fix):
                body_fixed += 1
                break

    # 2) Ensure "2.6 文献与标准依据" is a level-2 heading (outlineLvl=1) to appear in TOC when updated.
    outline_fixed = 0
    for p in doc.paragraphs:
        if (p.text or "").strip() == "2.6  文献与标准依据":
            _set_outline_level(p, 1)
            outline_fixed += 1
            break

    # 2.1) Make 2.6 show in the existing TOC block (static text) without requiring a Word "Update Field".
    # We insert a copied TOC line after "2.5  开发与运行环境\t8" if missing.
    toc_26_inserted = 0
    has_26 = any((p.text or "").strip().startswith("2.6  文献与标准依据\t") for p in doc.paragraphs)
    if not has_26:
        for p in doc.paragraphs:
            if (p.text or "").strip().startswith("2.5  开发与运行环境\t"):
                new_p = deepcopy(p._p)
                p._p.addnext(new_p)
                ts = new_p.xpath(".//w:t")
                if len(ts) >= 2:
                    ts[0].text = "2.6  "
                    ts[1].text = "文献与标准依据"
                    # Keep existing page number (usually 8) as-is.
                    toc_26_inserted = 1
                break

    # 3) Expand abstracts to meet typical length requirements without disturbing existing English-font runs.
    cn_abs_appended = 0
    en_abs_appended = 0

    cn_extra = (
        "在数据层方面，设计用户、部门、资产、拍卖、竞价、交易与处置等核心实体及其关联关系，"
        "保证业务数据一致性。针对竞价延时、成交确认超时与竞拍惩罚等业务规则，"
        "采用定时任务驱动状态变更并记录历史操作，确保过程可审计。"
    )
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip() == "本科毕业设计（论文）中文摘要":
            # Next non-empty paragraph is the abstract text.
            for q in doc.paragraphs[i + 1 :]:
                if (q.text or "").strip() and not (q.text or "").strip().startswith("关键词"):
                    if "在数据层方面" not in (q.text or ""):
                        q.add_run(cn_extra)
                        cn_abs_appended = 1
                    break
            break

    en_extra = (
        " The backend exposes RESTful APIs and persists core entities (user, department, asset, auction, bid, "
        "transaction, and disposal) in MySQL. Security is enforced through JWT-based authentication and role-based "
        "authorization to prevent unauthorized operations across the lifecycle. To improve auction fairness, the "
        "system supports dynamic extension of the end time when bids arrive near closing, and scheduled jobs handle "
        "auction settlement and transaction-confirmation timeouts automatically. Test results indicate that the main "
        "business chain can be completed reliably with clear status transitions and auditable records; under moderate "
        "concurrent access, key query and bidding endpoints maintain acceptable latency after index tuning. Overall, "
        "the platform is suitable for campus projects and enterprise intranet pilots, and it can be further enhanced "
        "with message queues, distributed locking, and finer-grained audit logging."
    )
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip() == "Abstract":
            for q in doc.paragraphs[i + 1 :]:
                t = (q.text or "").strip()
                if not t:
                    continue
                if t.startswith("Keywords"):
                    break
                if "The backend exposes RESTful APIs" not in t:
                    q.runs[0].text = (q.runs[0].text or "") + en_extra if q.runs else (t + en_extra)
                    en_abs_appended = 1
                break
            break

    # 4) Align Appendix A endpoints with actual controller mappings in this repo.
    appendix_fixed = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        replacement = APPENDIX_ENDPOINT_REPLACEMENTS.get(t)
        if not replacement:
            continue
        if p.runs:
            p.runs[0].text = replacement
            # Clear any remaining runs to avoid mixed old/new fragments.
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.text = replacement
        appendix_fixed += 1

    doc.core_properties.comments = (
        (doc.core_properties.comments or "")
        + f"\n[auto-fix] CH5 toc:{toc_fixed} body:{body_fixed}; outline2.6:{outline_fixed}; toc2.6+:{toc_26_inserted}; abs_cn+:{cn_abs_appended} abs_en+:{en_abs_appended}; appendix_ep:{appendix_fixed}"
    ).strip()
    doc.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Fix thesis numbering/titles in docx while preserving TOC structure.")
    parser.add_argument("input", help="Input .docx path")
    parser.add_argument("output", help="Output .docx path")
    args = parser.parse_args()

    fix_docx(args.input, args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
